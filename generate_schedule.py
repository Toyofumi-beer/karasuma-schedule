#!/usr/bin/env python3
"""
進学館烏丸御池校 月次スケジュール自動生成スクリプト
使い方: python3 generate_schedule.py
"""

import json
import calendar
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

WEEKDAY_JA = ["月", "火", "水", "木", "金", "土", "日"]

FONT_BOLD = "游ゴシック"
FONT_REG  = "游ゴシック"
FULLWIDTH_DIGITS = str.maketrans("0123456789", "０１２３４５６７８９")

# 週番号: 木曜始まり（木〜水）。5月第1木曜=第10週を基準に計算。
# ref_yearは学年度の4〜12月が属する年（例: 2026年5月→ref_year=2026）
def get_academic_start(ref_year):
    may1 = date(ref_year, 5, 1)
    thu_offset = (3 - may1.weekday()) % 7
    first_thu_may = may1 + timedelta(days=thu_offset)
    return first_thu_may - timedelta(weeks=9)  # 第1週の木曜（学年度起点）


def compute_week_map(no_class_periods, ref_year, last_day):
    """
    各日付の週番号を返す dict を生成。
    ルール: 日曜を除く全日をカウント。ただし授業なし期間(no_class_periods)はカウントしない。
    6カウントごとに週番号が1増える（通常の木〜水6授業日 = 1週）。
    固定祝日・年間イベント日はカウントされるが表示は別途制御する。
    """
    academic_start = get_academic_start(ref_year)
    count = 0
    week_map = {}
    d = academic_start
    while d <= last_day:
        if d.weekday() != 6:  # 日曜以外はカウント対象
            if not is_no_class_date(d, no_class_periods, ref_year):
                count += 1
        week_map[d] = (count - 1) // 6 + 1 if count > 0 else 1
        d += timedelta(days=1)
    return week_map


def load_annual_schedule():
    path = os.path.join(os.path.dirname(__file__), "annual_schedule.json")
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def normalize_date_key(month, day):
    return f"{month}/{day}"


def get_events_for_date(d, data, extra_events):
    """指定日の全イベントリストを返す。各要素は {"text": str, "color": str, "source": str}"""
    events = []
    key = normalize_date_key(d.month, d.day)
    weekday = WEEKDAY_JA[d.weekday()]

    # 年間固定イベント
    for event_name, ev in data["annual_events"].items():
        dates = ev.get("dates", [])
        if key in dates:
            color = ev.get("color", "")
            if "display_lines" in ev:
                # 複数行イベント（オープン模試・入試実践演習など）
                for line in ev["display_lines"]:
                    events.append({"text": line, "color": color, "source": "annual"})
            elif ev.get("display_format") == "monthly":
                # 月名入り表記（ミニ適性検査など）
                short_name = ev.get("short_name", event_name)
                session_months = ev.get("session_months", [])
                dates_list = ev.get("dates", [])
                try:
                    idx = dates_list.index(key)
                    session_month = session_months[idx]
                except (ValueError, IndexError):
                    session_month = d.month
                time = ev.get("time", "")
                month_fw = str(session_month).translate(FULLWIDTH_DIGITS)
                entry = f"{short_name} {month_fw}月授業({time})" if time else f"{short_name} {month_fw}月授業"
                events.append({"text": entry, "color": color, "source": "annual"})
            else:
                time = ev.get("time", "")
                note = ev.get("note", "")
                entry = f"{event_name}({time})" if time else event_name
                if note:
                    entry += f" {note}"
                events.append({"text": entry, "color": color, "source": "annual"})

    # 追加イベント（ユーザー入力）
    for ev in extra_events:
        if ev["date"] == key:
            events.append({"text": ev["text"], "color": ev.get("color", ""), "source": "extra"})

    # 週間固定授業
    if weekday in data["weekly_schedule"]:
        for lesson in data["weekly_schedule"][weekday]:
            grade = lesson.get("grade", "")
            name = lesson["name"]
            time = lesson["time"]
            entry = f"{grade}{name}({time})" if grade else f"{name}({time})"
            events.append({"text": entry, "color": "", "source": "weekly", "group": lesson.get("group")})

    return events


def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph_to_cell(cell, text, font_size=9, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = FONT_BOLD if bold else FONT_REG
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def is_no_class_date(d, no_class_periods, ref_year):
    """授業なし期間かどうか判定"""
    for period in no_class_periods:
        try:
            fm, fd = map(int, period["from"].split("/"))
            tm, td = map(int, period["to"].split("/"))
            start = date(ref_year if fm >= 4 else ref_year + 1, fm, fd)
            end   = date(ref_year if tm >= 4 else ref_year + 1, tm, td)
            if start <= d <= end:
                return True
        except Exception:
            pass
    return False


def get_no_class_short_label(d, no_class_periods, ref_year):
    """授業なし期間の短縮ラベル（先頭1文字）を返す。例: 夏期講習→夏"""
    for period in no_class_periods:
        try:
            fm, fd = map(int, period["from"].split("/"))
            tm, td = map(int, period["to"].split("/"))
            start = date(ref_year if fm >= 4 else ref_year + 1, fm, fd)
            end   = date(ref_year if tm >= 4 else ref_year + 1, tm, td)
            if start <= d <= end:
                label = period.get("label", "")
                return label[0] if label else ""
        except Exception:
            pass
    return ""


def _add_floating_image(paragraph, image_path, width_cm, x_cm, dy_cm=0.0):
    """段落にアンカーされた前面フローティング画像を追加する。
    x_cm: ページ左端からの水平位置 / dy_cm: アンカー段落からの縦オフセット。
    Word上では通常の図として選択・ドラッグで微調整できる。"""
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    inline = run._r.xpath('.//wp:inline')[0]

    anchor = OxmlElement('wp:anchor')
    for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                 ("simplePos", "0"), ("relativeHeight", "251658240"),
                 ("behindDoc", "0"), ("locked", "0"),
                 ("layoutInCell", "1"), ("allowOverlap", "1")):
        anchor.set(k, v)

    sp = OxmlElement('wp:simplePos'); sp.set('x', '0'); sp.set('y', '0')
    posH = OxmlElement('wp:positionH'); posH.set('relativeFrom', 'page')
    offH = OxmlElement('wp:posOffset'); offH.text = str(int(x_cm * 360000)); posH.append(offH)
    posV = OxmlElement('wp:positionV'); posV.set('relativeFrom', 'paragraph')
    offV = OxmlElement('wp:posOffset'); offV.text = str(int(dy_cm * 360000)); posV.append(offV)
    anchor.append(sp); anchor.append(posH); anchor.append(posV)

    for tag in ('wp:extent', 'wp:effectExtent'):
        el = inline.find(qn(tag))
        if el is not None:
            anchor.append(el)
    anchor.append(OxmlElement('wp:wrapNone'))
    for tag in ('wp:docPr', 'wp:cNvGraphicFramePr'):
        el = inline.find(qn(tag))
        if el is not None:
            anchor.append(el)
    anchor.append(inline.find(qn('a:graphic')))
    inline.getparent().replace(inline, anchor)


def generate_schedule(start_year, start_month, end_year, end_month,
                      extra_events=None, special_notices=None,
                      no_class_periods=None, output_path=None,
                      start_date=None, end_date=None, illustrations=None):
    """
    月次スケジュールWordファイルを生成する。

    extra_events: [{"date": "5/6", "text": "イベント名（時間）"}, ...]
    special_notices: [{"title": "□5/23 保護者会", "body": "内容..."}, ...]
    no_class_periods: [{"from": "7/17", "to": "8/31", "label": "夏期講習"}, ...]
    start_date / end_date: date オブジェクトで渡すと日付単位で範囲を指定できる
    """
    if extra_events is None:
        extra_events = []
    if special_notices is None:
        special_notices = []
    if no_class_periods is None:
        no_class_periods = []

    if start_date is not None:
        start_year, start_month = start_date.year, start_date.month
    if end_date is None:
        end_date = date(end_year, end_month, calendar.monthrange(end_year, end_month)[1])
    else:
        end_year, end_month = end_date.year, end_date.month
    if start_date is None:
        start_date = date(start_year, start_month, 1)

    # 学年度基準年（4月以降を基準）
    ref_year = start_year if start_month >= 4 else start_year - 1

    data = load_annual_schedule()

    doc = Document()

    # ページ余白を狭く
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # タイトル
    title_text = f"◆{data['school_name']} {start_month}〜{end_month}月スケジュール◆"
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text)
    title_run.font.size = Pt(13)
    title_run.font.bold = True
    title_run.font.name = FONT_BOLD
    title_para.paragraph_format.space_after = Pt(6)

    # スケジュール表（Excel版と同じ4列デザイン）
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 列幅設定
    col_widths = [Cm(1.8), Cm(1.0), Cm(1.0), Cm(13.6)]
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # ヘッダー行
    hdr = table.rows[0].cells
    for i, text in enumerate(["日程", "曜", "週", "集団授業、テスト、イベント"]):
        hdr[i].paragraphs[0].clear()
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = FONT_BOLD
        set_cell_background(hdr[i], "CCFFFF")

    # 日付ループ
    cur = start_date
    last_day = end_date

    week_map = compute_week_map(no_class_periods, ref_year, last_day)

    # カリキュラム月開始: (week_num-1)%3==0 の木曜に表示（5/28=第13週→６月授業開始 など）
    shown_curriculum_months = set()
    CURRICULUM_MONTHS = [2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 1]  # 2月始まり順

    event_cells_by_date = {}  # イラスト配置用: "7/7" -> イベントセル

    while cur <= last_day:
        weekday = WEEKDAY_JA[cur.weekday()]
        week_num = week_map.get(cur, 1)
        events = get_events_for_date(cur, data, extra_events)

        # 行追加
        row = table.add_row()
        cells = row.cells

        # 各種フラグ
        is_sunday = cur.weekday() == 6
        key = normalize_date_key(cur.month, cur.day)
        is_holiday = any(ev["date"] == key and "休館" in ev["text"] for ev in extra_events)
        no_class = is_no_class_date(cur, no_class_periods, ref_year)
        is_summer_closed = no_class and cur.weekday() in (3, 6)  # 授業なし期間の木・日は休館
        has_open_moshi = any("オープン模試" in ev["text"] for ev in events)
        has_jissen = any("入試実践演習" in ev["text"] for ev in events)
        has_no_class_event = any("平常授業なし" in ev["text"] for ev in events if ev["source"] != "weekly")
        has_gion_obon = any(
            ("祇園祭" in ev["text"] or "お盆" in ev["text"])
            for ev in events if ev["source"] == "extra"
        )

        # 授業なし期間・オープン模試・入試実践演習・平常授業なし日は週固定授業を非表示
        # （夏期講習中は欠席サポート教室も実施しないため、ここでまとめて消える）
        if no_class or has_open_moshi or has_jissen or has_no_class_event or has_gion_obon:
            events = [ev for ev in events if ev["source"] != "weekly"]

        # 土曜は「授業欠席サポート教室」を最上段に再配置（通常期のみ）
        support_event = next((ev for ev in events if "欠席サポート" in ev["text"]), None)
        if weekday == "土" and support_event is not None:
            events = [ev for ev in events if "欠席サポート" not in ev["text"]]
            events.insert(0, dict(support_event, source="support", group=None))

        # カリキュラム月開始の記載（(week_num-1)%3==0 の木曜。ただし平常授業なし日は表示しない）
        month_start_text = ""
        if weekday == "木" and (week_num - 1) % 3 == 0 and not has_no_class_event and not no_class:
            curriculum_month_num = (week_num - 1) // 3  # 0始まり (0=Feb, 4=Jun...)
            curriculum_cal_month = CURRICULUM_MONTHS[curriculum_month_num % 12]
            if curriculum_cal_month not in shown_curriculum_months:
                month_fw = str(curriculum_cal_month).translate(FULLWIDTH_DIGITS)
                month_start_text = f"【{month_fw}月授業開始】\n"
                shown_curriculum_months.add(curriculum_cal_month)

        # 行の背景色と文字色を決定（Excel版と同じ）
        if has_open_moshi:
            row_bg, row_fg = "FFFF00", None              # 黄背景・黒文字
        elif has_gion_obon:
            row_bg, row_fg = "7030A0", (255, 255, 255)   # 紫背景・白文字（祇園祭・お盆）
        elif has_jissen or is_sunday or is_holiday or is_summer_closed:
            row_bg, row_fg = "000000", (255, 255, 255)   # 黒背景・白文字
        elif no_class:
            row_bg, row_fg = "BDD7EE", None              # 水色（夏期講習通常日）
        else:
            row_bg, row_fg = None, None

        if row_bg:
            for c in cells:
                set_cell_background(c, row_bg)

        # 表示行リスト構築（Excel版と同じ: 追加/年間 → 週固定はgroupで横並び）
        all_lines = []
        if month_start_text:
            for t in month_start_text.strip().split("\n"):
                if t.strip():
                    all_lines.append({"text": t, "color": "", "source": "header"})
        for ev in [e for e in events if e.get("source") != "weekly"]:
            if ev["text"].strip():
                all_lines.append(ev)
        _seen_g = set()
        for ev in [e for e in events if e.get("source") == "weekly"]:
            if not ev["text"].strip():
                continue
            g = ev.get("group")
            if g is None:
                all_lines.append(ev)
            elif g not in _seen_g:
                _seen_g.add(g)
                _gt = [e["text"] for e in events
                       if e.get("source") == "weekly" and e.get("group") == g and e["text"].strip()]
                all_lines.append({"text": "　".join(_gt), "color": "", "source": "weekly"})

        # 夏期休館日（木・日）は「休館」のみ表示（祇園祭・お盆表示を除く）
        if is_summer_closed and not has_gion_obon:
            all_lines = [{"text": "休館", "color": "", "source": "closed"}]

        # 週表示（週番号はweek_overrides優先。PDF年間スケジュール準拠）
        week_override = data.get("week_overrides", {}).get(key)
        if is_sunday or has_open_moshi or has_jissen or is_holiday or is_summer_closed or has_gion_obon:
            week_display = ""
        elif no_class:
            week_display = get_no_class_short_label(cur, no_class_periods, ref_year)
        elif week_override is not None:
            week_display = str(week_override)
        else:
            week_display = str(week_num)

        # 日付・曜・週セル
        for idx, val in ((0, f"{cur.month}/{cur.day}"), (1, weekday), (2, week_display)):
            c = cells[idx]
            c.paragraphs[0].clear()
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.name = FONT_BOLD
            if row_fg:
                r.font.color.rgb = RGBColor(*row_fg)

        # イベントセル
        event_cell = cells[3]
        event_cell.paragraphs[0].clear()
        event_cells_by_date[key] = event_cell
        first = True
        for item in all_lines:
            text = item["text"]
            if not text.strip():
                continue
            if first:
                p3 = event_cell.paragraphs[0]
                first = False
            else:
                p3 = event_cell.add_paragraph()
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(0)
            # 休館・模試・演習は中央寄せ
            if "休館" in text or has_open_moshi or has_jissen:
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run(text)
            run3.font.size = Pt(9)
            run3.font.bold = True
            run3.font.name = FONT_BOLD
            # 文字色の優先順位: 行色 > イベント個別色 > ★
            if row_fg:
                run3.font.color.rgb = RGBColor(*row_fg)
            elif item.get("color") == "red" or "適性検査" in text or "総合回テスト" in text:
                run3.font.color.rgb = RGBColor(192, 0, 0)
            elif item.get("color") == "purple":
                run3.font.color.rgb = RGBColor(112, 48, 160)
            elif text.startswith("★"):
                run3.font.color.rgb = RGBColor(192, 0, 0)

        cur += timedelta(days=1)

    # 列幅を固定レイアウトで確定（行追加後に全セルへ適用しないとWordで幅が崩れる。
    # tblGridにも直接設定しないとPagesで列幅が無視される）
    tblPr = table._tbl.tblPr
    layout_el = OxmlElement('w:tblLayout')
    layout_el.set(qn('w:type'), 'fixed')
    tblPr.append(layout_el)
    for gc, width in zip(table._tbl.tblGrid.findall(qn('w:gridCol')), col_widths):
        gc.set(qn('w:w'), str(width.twips))
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # イラスト配置（日付の行にアンカーした前面フローティング画像。Word上でドラッグ調整可）
    if illustrations:
        for il in illustrations:
            img = il.get("image", "")
            if not img or not os.path.exists(img):
                continue
            d = (il.get("date") or "").strip()
            para = (event_cells_by_date[d].paragraphs[0]
                    if d in event_cells_by_date else title_para)
            _add_floating_image(
                para, img,
                width_cm=float(il.get("width_cm", 3.0)),
                x_cm=float(il.get("x_cm", 15.8)),
                dy_cm=float(il.get("dy_cm", 0.0)),
            )

    doc.add_paragraph()

    # 備考・お知らせセクション
    for notice in data["fixed_notices"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(notice)
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.name = FONT_BOLD

    # 月固有のお知らせ
    for sn in special_notices:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(sn["title"])
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = FONT_BOLD
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        run2 = p2.add_run(sn["body"])
        run2.font.size = Pt(8.5)
        run2.font.bold = True
        run2.font.name = FONT_BOLD

    # 標準注意事項
    doc.add_paragraph()
    notice_header = doc.add_paragraph()
    run = notice_header.add_run("＜事務局よりお知らせ＞")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.name = FONT_BOLD
    notice_header.paragraph_format.space_after = Pt(2)

    for sn in data["standard_notices"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(sn)
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.name = FONT_BOLD

    # 保存
    if output_path is None:
        output_path = f"{start_year}karasuma_{start_month:02d}-{end_month:02d}.docx"
    output_path = os.path.join(os.path.dirname(__file__), output_path)
    doc.save(output_path)
    print(f"✅ 保存しました: {output_path}")
    return output_path


def _convert_inline_to_shared(path):
    """openpyxlが書くインライン文字列をsharedStrings形式（Excelネイティブ）に変換する。
    インライン形式のままだとQuick Look・Numbers等で文字色（部分赤字）が表示されないため。"""
    import zipfile
    import re as _re
    with zipfile.ZipFile(path) as z:
        contents = {n: z.read(n) for n in z.namelist()}
    sheets = [n for n in contents if _re.match(r"xl/worksheets/sheet\d+\.xml$", n)]
    sst = []

    def _repl(m):
        pre, inner = m.group(1), m.group(2)
        inner = inner.replace("<t>", '<t xml:space="preserve">')
        sst.append("<si>" + inner + "</si>")
        return pre.replace('t="inlineStr"', 't="s"') + f"<v>{len(sst) - 1}</v></c>"

    changed = False
    for n in sheets:
        x = contents[n].decode("utf-8")
        new = _re.sub(r'(<c [^>]*t="inlineStr"[^>]*>)<is>(.*?)</is></c>', _repl, x, flags=_re.S)
        if new != x:
            contents[n] = new.encode("utf-8")
            changed = True
    if not changed:
        return

    contents["xl/sharedStrings.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"'
        f' count="{len(sst)}" uniqueCount="{len(sst)}">{"".join(sst)}</sst>'
    ).encode("utf-8")

    ct = contents["[Content_Types].xml"].decode("utf-8")
    if "sharedStrings" not in ct:
        ct = ct.replace("</Types>",
            '<Override PartName="/xl/sharedStrings.xml" ContentType="application/'
            'vnd.openxmlformats-officedocument.spreadsheetml.sharedStrings+xml"/></Types>')
        contents["[Content_Types].xml"] = ct.encode("utf-8")

    wr = contents["xl/_rels/workbook.xml.rels"].decode("utf-8")
    if "sharedStrings" not in wr:
        ids = [int(i) for i in _re.findall(r'Id="rId(\d+)"', wr)]
        rid = (max(ids) + 1) if ids else 1
        wr = wr.replace("</Relationships>",
            f'<Relationship Id="rId{rid}" Type="http://schemas.openxmlformats.org/'
            'officeDocument/2006/relationships/sharedStrings" Target="sharedStrings.xml"/>'
            "</Relationships>")
        contents["xl/_rels/workbook.xml.rels"] = wr.encode("utf-8")

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for n, data in contents.items():
            z.writestr(n, data)


def generate_schedule_excel(start_year, start_month, end_year, end_month,
                             extra_events=None, special_notices=None,
                             no_class_periods=None, output_path=None,
                             start_date=None, end_date=None):
    """
    月次スケジュールをExcelファイルで生成。
    通信のスケジュール.xlsx と同じデザイン（游ゴシック・thin罫線・CCFFFF ヘッダー）を踏襲。
    start_date / end_date: date オブジェクトで渡すと日付単位で範囲を指定できる
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

    if extra_events is None:    extra_events = []
    if special_notices is None: special_notices = []
    if no_class_periods is None: no_class_periods = []

    if start_date is not None:
        start_year, start_month = start_date.year, start_date.month
    if end_date is None:
        end_date = date(end_year, end_month, calendar.monthrange(end_year, end_month)[1])
    else:
        end_year, end_month = end_date.year, end_date.month
    if start_date is None:
        start_date = date(start_year, start_month, 1)

    ref_year = start_year if start_month >= 4 else start_year - 1
    data = load_annual_schedule()

    # ── スタイル定義 ────────────────────────────────────────
    thin = Side(style="thin")
    BD = Border(left=thin, right=thin, top=thin, bottom=thin)

    def fnt(color="000000", size=11):
        argb = color if len(color) == 8 else "FF" + color
        return Font(name="游ゴシック", size=size, bold=True, color=argb)

    def fill(hex_color):
        argb = hex_color if len(hex_color) == 8 else "FF" + hex_color
        return PatternFill(fill_type="solid", fgColor=argb)

    AL_C  = Alignment(horizontal="center", vertical="center")
    AL_CT = Alignment(horizontal="center", vertical="top")
    AL_LT = Alignment(horizontal="left",   vertical="top", wrap_text=True)

    wb = Workbook()
    ws = wb.active
    ws.title = f"{start_year}年{start_month}〜{end_month}月"

    # 列幅（テンプレートに合わせる）
    ws.column_dimensions["A"].width = 3.25   # 余白
    ws.column_dimensions["B"].width = 7.5    # 日付
    ws.column_dimensions["C"].width = 3.25   # 曜日
    ws.column_dimensions["D"].width = 5.0    # 週
    ws.column_dimensions["E"].width = 65.0   # 授業内容

    ri = 1  # 行インデックス

    # タイトル行
    ws.merge_cells(f"B{ri}:E{ri}")
    c = ws[f"B{ri}"]
    c.value = f"◆{data['school_name']} {start_month}〜{end_month}月スケジュール◆"
    c.font  = Font(name="游ゴシック", size=14, bold=True)
    c.alignment = AL_C
    ws.row_dimensions[ri].height = 24
    ri += 1

    # ヘッダー行
    for col, val in [("B","日程"),("C","曜"),("D","週"),("E","集団授業、テスト、イベント")]:
        c = ws[f"{col}{ri}"]
        c.value = val; c.font = fnt(); c.border = BD
        c.fill = fill("CCFFFF"); c.alignment = AL_C
    ws.row_dimensions[ri].height = 22
    ri += 1

    # ── データ行 ────────────────────────────────────────────
    shown_curriculum_months = set()
    CURRICULUM_MONTHS = [2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 1]
    cur      = start_date
    last_day = end_date

    week_map = compute_week_map(no_class_periods, ref_year, last_day)

    while cur <= last_day:
        weekday  = WEEKDAY_JA[cur.weekday()]
        week_num = week_map.get(cur, 1)
        events   = get_events_for_date(cur, data, extra_events)
        key      = normalize_date_key(cur.month, cur.day)

        is_sunday        = cur.weekday() == 6
        is_holiday       = any(ev["date"] == key and "休館" in ev["text"] for ev in extra_events)
        no_class         = is_no_class_date(cur, no_class_periods, ref_year)
        is_summer_closed = no_class and cur.weekday() in (3, 6)  # 授業なし期間の木・日は休館
        has_open_moshi   = any("オープン模試"    in ev["text"] for ev in events)
        has_jissen       = any("入試実践演習"    in ev["text"] for ev in events)
        has_no_class_ev  = any("平常授業なし"    in ev["text"] for ev in events if ev["source"] != "weekly")
        has_gion_obon    = any(("祇園祭" in ev["text"] or "お盆" in ev["text"])
                               for ev in events if ev["source"] == "extra")

        if no_class or has_open_moshi or has_jissen or has_no_class_ev or has_gion_obon:
            events = [ev for ev in events if ev["source"] != "weekly"]

        # 土曜は「授業欠席サポート教室」を最上段に再配置（通常期のみ。夏期講習中は実施なし）
        support = next((ev for ev in events if "欠席サポート" in ev["text"]), None)
        if weekday == "土" and support is not None:
            events = [ev for ev in events if "欠席サポート" not in ev["text"]]
            events.insert(0, dict(support, source="support", group=None))

        # カリキュラム月ラベル
        month_start_text = ""
        if weekday == "木" and (week_num - 1) % 3 == 0 and not has_no_class_ev and not no_class:
            cm_num = (week_num - 1) // 3
            cm_cal = CURRICULUM_MONTHS[cm_num % 12]
            if cm_cal not in shown_curriculum_months:
                month_fw = str(cm_cal).translate(FULLWIDTH_DIGITS)
                month_start_text = f"【{month_fw}月授業開始】"
                shown_curriculum_months.add(cm_cal)

        # 週表示（オープン模試・入試実践演習・休館日・夏期休館日・祇園祭お盆は空白）
        wk_override = data.get("week_overrides", {}).get(normalize_date_key(cur.month, cur.day))
        if is_sunday or has_open_moshi or has_jissen or is_holiday or is_summer_closed or has_gion_obon:
            wk_disp = ""
        elif no_class:
            wk_disp = get_no_class_short_label(cur, no_class_periods, ref_year)
        elif wk_override is not None:
            wk_disp = str(wk_override)
        else:
            wk_disp = str(week_num)

        # イベント行リスト構築（週次イベントはgroupで横並び）
        all_lines = []
        if month_start_text:
            all_lines.append((month_start_text, "000000"))
        for ev in [e for e in events if e.get("source") != "weekly"]:
            if ev["text"].strip():
                is_red = ev.get("color") == "red" or "適性検査" in ev["text"] or "総合回テスト" in ev["text"]
                all_lines.append((ev["text"], "C00000" if is_red else "000000"))
        _seen_g = set()
        for ev in [e for e in events if e.get("source") == "weekly"]:
            if not ev["text"].strip():
                continue
            g = ev.get("group")
            if g is None:
                all_lines.append((ev["text"], "000000"))
            elif g not in _seen_g:
                _seen_g.add(g)
                _gt = [e["text"] for e in events
                       if e.get("source") == "weekly" and e.get("group") == g and e["text"].strip()]
                all_lines.append(("　".join(_gt), "000000"))
        ev_text = "\n".join(t for t, _ in all_lines)

        # 夏期休館日（木・日）は内容を「休館」に上書き（祇園祭・お盆を除く）
        if is_summer_closed and not has_gion_obon:
            all_lines = [("休館", "FFFFFF")]
            ev_text   = "休館"

        # 行背景・文字色
        if has_open_moshi:
            bg, fg = "FFFF00", "000000"
        elif has_gion_obon:
            bg, fg = "7030A0", "FFFFFF"   # 紫（祇園祭・お盆）
        elif has_jissen or is_sunday or is_holiday or is_summer_closed:
            bg, fg = "000000", "FFFFFF"   # 黒（休館）
        elif no_class:
            bg, fg = "BDD7EE", "000000"   # 水色（夏期講習通常日）
        else:
            bg, fg = None, "000000"

        # 日付・曜日・週セル
        for col, val, align in [
            ("B", f"{cur.month}/{cur.day}", AL_CT),
            ("C", weekday,  AL_CT),
            ("D", wk_disp,  AL_CT),
        ]:
            c = ws[f"{col}{ri}"]
            c.value     = val
            c.border    = BD
            c.alignment = align
            c.fill      = fill(bg) if bg else PatternFill(fill_type=None)
            c.font      = fnt(color=fg)

        # イベントセル — 行背景がある場合は単色、なければ行ごとに色を分ける
        ev_cell = ws[f"E{ri}"]
        ev_cell.border    = BD
        # 休館表示は中央寄せ
        if "休館" in ev_text:
            ev_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        else:
            ev_cell.alignment = AL_LT
        ev_cell.fill      = fill(bg) if bg else PatternFill(fill_type=None)

        if bg and fg != "000000":
            # 白文字行（休館・祇園祭お盆など）→ セル全体を前景色で統一
            ev_cell.value = ev_text
            ev_cell.font  = fnt(color=fg)
        else:
            # 背景色なし → RichText で行ごとに色指定
            from openpyxl.cell.rich_text import CellRichText, TextBlock
            from openpyxl.cell.text import InlineFont

            # all_lines は上で構築済み
            if not all_lines:
                ev_cell.value = ""
                ev_cell.font  = fnt()
            elif all(c == all_lines[0][1] for _, c in all_lines):
                # 全行同じ色ならシンプルに設定
                ev_cell.value = "\n".join(t for t, _ in all_lines)
                ev_cell.font  = fnt(color=all_lines[0][1])
            else:
                # 混在 → RichText
                def ifnt(color):
                    # alpha FF = 完全不透明
                    argb = color if len(color) == 8 else "FF" + color
                    return InlineFont(rFont="游ゴシック", sz=11, b=True, color=argb)

                blocks = []
                for i, (text, color) in enumerate(all_lines):
                    if i > 0:
                        blocks.append(TextBlock(ifnt("000000"), "\n"))
                    blocks.append(TextBlock(ifnt(color), text))
                ev_cell.value = CellRichText(*blocks)
                ev_cell.font  = fnt()  # ベースフォントを明示（游ゴシック太字）

        # 行高さ: 折り返しを考慮した行数 × 15 + パディング8
        def _est_lines(text, col_w=45):
            if not text or not text.strip():
                return 1
            total = 0
            for seg in text.split("\n"):
                w = sum(2 if ord(c) > 127 else 1 for c in seg)
                total += max(1, -(-w // col_w))  # 切り上げ除算
            return max(1, total)
        n_lines = _est_lines(ev_text)
        ws.row_dimensions[ri].height = max(21, n_lines * 15 + 8)
        ri  += 1
        cur += timedelta(days=1)

    # 保存
    if output_path is None:
        output_path = os.path.join(os.path.dirname(__file__),
                                    f"{start_year}karasuma_{start_month:02d}-{end_month:02d}.xlsx")
    else:
        output_path = os.path.join(os.path.dirname(__file__), output_path)
    wb.save(output_path)
    _convert_inline_to_shared(output_path)
    print(f"✅ Excel保存しました: {output_path}")
    return output_path


def interactive_mode():
    """対話モードで月次スケジュールを作成"""
    print("=" * 50)
    print("進学館烏丸御池校 スケジュール作成ツール")
    print("=" * 50)

    # 対象月の入力
    print("\n【対象期間】")
    start_input = input("開始月を入力してください (例: 2026/7): ").strip()
    end_input = input("終了月を入力してください (例: 2026/8): ").strip()

    try:
        sy, sm = map(int, start_input.split("/"))
        ey, em = map(int, end_input.split("/"))
    except ValueError:
        print("❌ 入力形式が正しくありません。YYYY/M の形式で入力してください。")
        return

    # 追加イベント
    print("\n【追加イベント】（年間固定以外の特別日程）")
    print("形式: 月/日 イベント内容 （例: 7/20 保護者会（19:30〜20:00））")
    print("入力が終わったら空のままEnterを押してください")
    extra_events = []
    while True:
        ev_input = input("  追加イベント: ").strip()
        if not ev_input:
            break
        parts = ev_input.split(" ", 1)
        if len(parts) == 2:
            extra_events.append({"date": parts[0], "text": parts[1]})
        else:
            print("  ⚠️ スペースで日付とイベント名を区切ってください")

    # 特別お知らせ
    print("\n【特別お知らせ】（保護者会・サポートデイの詳細説明など）")
    print("タイトルを入力し、次に本文を入力。複数追加可。空のままEnterで終了。")
    special_notices = []
    while True:
        title = input("  お知らせタイトル（例: □7/20 オンライン保護者会）: ").strip()
        if not title:
            break
        body = input("  お知らせ本文: ").strip()
        special_notices.append({"title": title, "body": body})

    # 出力ファイル名
    default_name = f"{sy}karasuma_{sm:02d}-{em:02d}.docx"
    fname = input(f"\n出力ファイル名 (デフォルト: {default_name}): ").strip()
    if not fname:
        fname = default_name

    print("\n生成中...")
    path = generate_schedule(sy, sm, ey, em, extra_events, special_notices, fname)
    print(f"\n完了！ファイルを確認してください: {path}")
    print("Wordで開いて確認後、PDFとしてエクスポートして配信してください。")


if __name__ == "__main__":
    interactive_mode()
